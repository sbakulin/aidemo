import os
import io
import boto3
from typing import Optional, List, BinaryIO
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
import litellm
from app.storage import storage
from app.models import ProcessingJobType, ProcessingJobStatus


class S3Service:
    def __init__(self):
        self.s3_client = None
        self.bucket_name = os.getenv("S3_BUCKET_NAME")
        
        if os.getenv("AWS_ACCESS_KEY_ID") and os.getenv("AWS_SECRET_ACCESS_KEY"):
            self.s3_client = boto3.client(
                's3',
                aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
                region_name=os.getenv("AWS_REGION", "us-east-1")
            )
    
    def upload_file(self, file_data: bytes, key: str, content_type: str = "application/octet-stream") -> Optional[str]:
        """Upload file to S3 and return the key"""
        if not self.s3_client or not self.bucket_name:
            return f"memory://{key}"
        
        try:
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=key,
                Body=file_data,
                ContentType=content_type
            )
            return key
        except Exception as e:
            print(f"Error uploading to S3: {e}")
            return None
    
    def get_file_url(self, key: str, expiration: int = 3600) -> Optional[str]:
        """Generate presigned URL for file access"""
        if key.startswith("memory://"):
            return key
        
        if not self.s3_client or not self.bucket_name:
            return None
        
        try:
            url = self.s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': self.bucket_name, 'Key': key},
                ExpiresIn=expiration
            )
            return url
        except Exception as e:
            print(f"Error generating presigned URL: {e}")
            return None
    
    def download_file(self, key: str) -> Optional[bytes]:
        """Download file from S3"""
        if key.startswith("memory://"):
            return None
        
        if not self.s3_client or not self.bucket_name:
            return None
        
        try:
            response = self.s3_client.get_object(Bucket=self.bucket_name, Key=key)
            return response['Body'].read()
        except Exception as e:
            print(f"Error downloading from S3: {e}")
            return None


class PDFService:
    def __init__(self, s3_service: S3Service):
        self.s3_service = s3_service
    
    def extract_text(self, pdf_data: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(pdf_data)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    async def analyze_pdf_visual(self, pdf_data: bytes) -> dict:
        """Analyze PDF pages using visual LLM"""
        try:
            
            return {
                "pages_analyzed": 0,
                "plots": [],
                "tables": [],
                "formulas": [],
                "note": "Visual analysis not yet implemented - requires vision model"
            }
        except Exception as e:
            print(f"Error analyzing PDF visually: {e}")
            return {"error": str(e)}
    
    async def process_pdf(self, article_id: str, pdf_s3_key: str) -> dict:
        """Process PDF: extract text and analyze visually"""
        job = storage.create_processing_job(ProcessingJobType.PDF)
        
        try:
            storage.update_processing_job(job.id, ProcessingJobStatus.PROCESSING)
            
            pdf_data = self.s3_service.download_file(pdf_s3_key)
            if not pdf_data:
                raise Exception("Could not download PDF from S3")
            
            text = self.extract_text(pdf_data)
            
            visual_analysis = await self.analyze_pdf_visual(pdf_data)
            
            storage.update_article(article_id, pdf_text=text)
            
            await self.generate_embeddings(article_id, text)
            
            result = {
                "text_length": len(text),
                "visual_analysis": visual_analysis
            }
            
            storage.update_processing_job(job.id, ProcessingJobStatus.COMPLETED, result=result)
            return {"job_id": job.id, "status": "completed", "result": result}
            
        except Exception as e:
            storage.update_processing_job(job.id, ProcessingJobStatus.FAILED, error=str(e))
            return {"job_id": job.id, "status": "failed", "error": str(e)}
    
    async def generate_embeddings(self, article_id: str, text: str):
        """Generate embeddings for text chunks"""
        try:
            chunk_size = 1000
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
            
            for idx, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue
                
                try:
                    response = await litellm.aembedding(
                        model="text-embedding-ada-002",
                        input=[chunk]
                    )
                    
                    embedding = response.data[0]['embedding']
                    key = f"article:{article_id}:chunk:{idx}"
                    storage.store_embedding(key, embedding)
                except Exception as e:
                    print(f"Error generating embedding for chunk {idx}: {e}")
                    
        except Exception as e:
            print(f"Error in generate_embeddings: {e}")


class STTService:
    async def transcribe_audio(self, audio_data: bytes, audio_format: str = "mp3") -> str:
        """Transcribe audio using Whisper via LiteLLM"""
        try:
            temp_file = f"/tmp/audio_{storage.generate_id()}.{audio_format}"
            with open(temp_file, "wb") as f:
                f.write(audio_data)
            
            response = await litellm.atranscription(
                model="whisper-1",
                file=open(temp_file, "rb")
            )
            
            os.remove(temp_file)
            
            return response.text
            
        except Exception as e:
            print(f"Error transcribing audio: {e}")
            return ""
    
    async def process_audio(self, message_id: str, audio_s3_key: str) -> dict:
        """Process audio: transcribe and update message"""
        job = storage.create_processing_job(ProcessingJobType.AUDIO)
        
        try:
            storage.update_processing_job(job.id, ProcessingJobStatus.PROCESSING)
            
            s3_service = S3Service()
            audio_data = s3_service.download_file(audio_s3_key)
            if not audio_data:
                raise Exception("Could not download audio from S3")
            
            transcription = await self.transcribe_audio(audio_data)
            
            storage.update_message(message_id, audio_transcription=transcription)
            
            result = {"transcription": transcription}
            storage.update_processing_job(job.id, ProcessingJobStatus.COMPLETED, result=result)
            
            return {"job_id": job.id, "status": "completed", "result": result}
            
        except Exception as e:
            storage.update_processing_job(job.id, ProcessingJobStatus.FAILED, error=str(e))
            return {"job_id": job.id, "status": "failed", "error": str(e)}


class RAGService:
    async def query(self, query: str, limit: int = 5) -> dict:
        """Query using RAG across all content"""
        try:
            response = await litellm.aembedding(
                model="text-embedding-ada-002",
                input=[query]
            )
            
            query_embedding = response.data[0]['embedding']
            
            similar = storage.search_similar(query_embedding, limit=limit)
            
            contexts = []
            for key, score in similar:
                parts = key.split(":")
                if parts[0] == "article":
                    article_id = parts[1]
                    article = storage.get_article(article_id)
                    if article:
                        contexts.append({
                            "type": "article",
                            "id": article_id,
                            "title": article.title,
                            "score": score
                        })
            
            return {
                "query": query,
                "contexts": contexts,
                "note": "RAG query completed - contexts retrieved"
            }
            
        except Exception as e:
            print(f"Error in RAG query: {e}")
            return {"error": str(e)}


class WordExportService:
    def __init__(self, s3_service: S3Service):
        self.s3_service = s3_service
    
    async def generate_document(self, title: str, dialog_id: Optional[str] = None,
                               article_ids: Optional[List[str]] = None) -> dict:
        """Generate Word document with findings"""
        job = storage.create_processing_job(ProcessingJobType.EXPORT)
        
        try:
            storage.update_processing_job(job.id, ProcessingJobStatus.PROCESSING)
            
            doc = Document()
            doc.add_heading(title, 0)
            
            if article_ids:
                doc.add_heading("Articles", level=1)
                for article_id in article_ids:
                    article = storage.get_article(article_id)
                    if article:
                        doc.add_heading(article.title, level=2)
                        if article.url:
                            doc.add_paragraph(f"URL: {article.url}")
                        if article.pdf_text:
                            doc.add_paragraph(article.pdf_text[:500] + "...")
                        
                        if article.comments:
                            doc.add_heading("Comments:", level=3)
                            for comment in article.comments:
                                if comment.text:
                                    doc.add_paragraph(comment.text)
                                if comment.audio_transcription:
                                    doc.add_paragraph(f"[Audio transcription]: {comment.audio_transcription}")
            
            if dialog_id:
                dialog = storage.get_dialog(dialog_id)
                if dialog:
                    doc.add_heading(f"Dialog: {dialog.title}", level=1)
                    messages = storage.list_messages(dialog_id)
                    
                    for message in messages:
                        if message.text:
                            doc.add_paragraph(message.text)
                        if message.audio_transcription:
                            doc.add_paragraph(f"[Audio]: {message.audio_transcription}")
                        
                        if message.image_s3_keys:
                            doc.add_paragraph(f"[{len(message.image_s3_keys)} image(s) attached]")
            
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            doc_key = f"exports/document_{job.id}.docx"
            s3_key = self.s3_service.upload_file(
                doc_bytes.read(),
                doc_key,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            if not s3_key:
                raise Exception("Failed to upload document to S3")
            
            download_url = self.s3_service.get_file_url(s3_key, expiration=86400)  # 24 hours
            
            result = {
                "document_key": s3_key,
                "download_url": download_url
            }
            
            storage.update_processing_job(job.id, ProcessingJobStatus.COMPLETED, result=result)
            return {"job_id": job.id, "status": "completed", "result": result}
            
        except Exception as e:
            storage.update_processing_job(job.id, ProcessingJobStatus.FAILED, error=str(e))
            return {"job_id": job.id, "status": "failed", "error": str(e)}


s3_service = S3Service()
pdf_service = PDFService(s3_service)
stt_service = STTService()
rag_service = RAGService()
word_export_service = WordExportService(s3_service)
